#!/usr/bin/env python3
"""
Japanese Immigration Form PDF Generator  
Handles passport data extraction and PDF form population in Japanese
"""

import json
import tempfile
import os
from datetime import datetime
from typing import Dict, List, Any
from pikepdf import Pdf
import pikepdf
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import black
from docx import Document
from PIL import Image
import base64
import google.generativeai as genai


class JapaneseFormGenerator:
    """Main class for generating Japanese immigration forms from passport data"""
    
    def __init__(self, pdf_config_path: str = "/Users/hugo/passport_project/web_app/pdf_config/visa_form_mapping.json"):
        self.pdf_config_path = pdf_config_path
        self.config = self.load_config()
        self.current_year = datetime.now().year
        # The API key is configured globally in app.py
        
    def load_config(self) -> Dict[str, Any]:
        """Load PDF field mappings"""
        with open(self.pdf_config_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def convert_to_japanese_calendar(self, iso_date: str) -> str:
        """Convert ISO date to YYYY年MM月DD日 format"""
        from datetime import datetime
        
        # Parse the date
        if iso_date and '-' in iso_date:
            dt = datetime.strptime(iso_date, '%Y-%m-%d')
            
            # Format: YYYY年MM月DD日
            return dt.strftime('%Y年%m月%d日')
        
        return iso_date

    def map_passport_to_form_fields(self, passport_data: Dict[str, Any]) -> Dict[str, str]:
        """Map passport data to PDF form fields"""
        form_values = {}
        
        # Japanese name formatting (surname first)
        surname = passport_data.get('surname', '')
        given_names = passport_data.get('given_names', '')
        form_values['full_name_japanese_order'] = f"{surname.upper()} {given_names.upper()}".strip()
        
        # Passport number
        form_values['passport_number'] = passport_data.get('passport_number', '')
        
        # Date conversion to Japanese format
        birth_date = passport_data.get('date_of_birth', '')
        if birth_date:
            form_values['date_of_birth'] = self.convert_to_japanese_calendar(birth_date)
        
        # Nationality
        nationality_en = passport_data.get('nationality', '')
        nationality_map = {
            'UNITED STATES': 'アメリカ合衆国',
            'JAPAN': '日本',
            'UNITED KINGDOM': 'イギリス',
            'CANADA': 'カナダ',
            'AUSTRALIA': 'オーストラリア',
            # Add more as needed
        }
        form_values['nationality'] = nationality_map.get(nationality_en.upper(), nationality_en)
        
        # Birth place
        place_of_birth = passport_data.get('place_of_birth', '')
        place_map = {
            'UNITED STATES': 'アメリカ合衆国',
            'CALIFORNIA': 'カリフォルニア州',
            'NEW YORK': 'ニューヨーク州',
            # Add more mappings
        }
        if place_of_birth.upper() in place_map:
            form_values['birth_place'] = place_map[place_of_birth.upper()]
        else:
            form_values['birth_place'] = place_of_birth
        
        # Issuing authority
        authority = passport_data.get('issuing_authority', '')
        form_values['issuing_authority'] = authority.upper() if authority else ''
        
        # Expiration date
        expiry_date = passport_data.get('date_of_expiration', '')
        if expiry_date:
            form_values['expiration_date'] = self.convert_to_japanese_calendar(expiry_date)
        
        return form_values

    def generate_pdf_form(self, member_data: Dict[str, Dict[str, str]], output_path: str) -> str:
        """Generate filled PDF form"""
        template_path = self.config.get('pdf_name', 'visa_request_form.pdf')
        full_template_path = os.path.join('template_file', template_path)
        
        if os.path.exists(full_template_path):
            # Use existing template and fill fields
            return self.fill_existing_pdf(full_template_path, member_data, output_path)
        else:
            # Generate new form based on template
            return self.generate_new_pdf(member_data, output_path)

    def fill_existing_pdf(self, template_path: str, data: Dict[str, Dict[str, str]], output_path: str) -> str:
        """Fill existing PDF template with data"""
        try:
            with Pdf.open(template_path) as pdf:
                # Get the form fields
                fields = {}
                for page in pdf.pages:
                    if '/Annots' in page:
                        for annot in page.Annots:
                            if '/T' in annot: # Field name
                                fields[str(annot.T)] = annot
                print(f"Found PDF fields: {list(fields.keys())}")

                # Iterate through the data and fill the corresponding fields
                for member_type, member_data in data.items():
                    # Handle primary applicant
                    if member_type == 'primary':
                        mapping_config = self.config.get('primary_applicant', {})
                    # Handle accompanying members dynamically (accompanying1 through accompanying7)
                    elif member_type.startswith('accompanying'):
                        mapping_config = self.config.get(member_type, {})
                        if not mapping_config:
                            print(f"Warning: No mapping configuration found for {member_type}")
                            continue
                    else:
                        continue # Skip unknown member types

                    for logical_field, extracted_value in member_data.items():
                        pdf_field_name = mapping_config.get(logical_field)
                        if pdf_field_name and pdf_field_name in fields:
                            fields[pdf_field_name].V = extracted_value
                            fields[pdf_field_name].AP = fields[pdf_field_name].V # Update appearance stream
                            print(f"Filled PDF field {pdf_field_name} for {member_type} with {extracted_value}")
                        elif pdf_field_name:
                            print(f"Warning: PDF field '{pdf_field_name}' (mapped from '{logical_field}') for {member_type} not found in PDF template.")
                        else:
                            print(f"Warning: Logical field '{logical_field}' for {member_type} not found in mapping configuration.")
                
                pdf.save(output_path)
                return output_path
        except Exception as e:
            print(f"Error filling PDF: {e}")
            # If an error occurs, we should not return a corrupted PDF.
            # Re-raise the exception or handle it appropriately.
            raise RuntimeError(f"Failed to fill PDF form: {e}")

    # The create_pdf_overlay function is removed as it's not suitable for filling existing forms.
    # The generate_new_pdf function is also removed as it's not used in this context.
    def generate_new_pdf(self, member_data: Dict[str, Dict[str, str]], output_path: str) -> str:
        raise NotImplementedError("Generating new PDFs from scratch is not implemented.")

    def generate_word_document(self, data: Dict[str, Dict[str, str]], output_path: str) -> str:
        """Generate Word document with visa form data mimicking PDF content"""
        doc = Document()

        doc.add_paragraph("重複プロファイル解消のお願い")
        doc.add_paragraph("ビザ面接の予約をしようとしたところ、以下のメッセージが表示されました。")
        doc.add_paragraph("「あなたの個人情報は、当社のデータベースに既に存在するプロファイルと一致します。」")
        doc.add_paragraph("しかし、古いプロファイルのログイン情報がわからず、現在予約手続きを進められない状況です。")
        doc.add_paragraph("お手数ですが、古いプロファイルを削除いただけますようお願い申し上げます。")

        # Primary Applicant
        doc.add_heading("主たる申請者", level=1)
        primary_data = data.get("primary", {})
        doc.add_paragraph(f"名前：{primary_data.get('full_name_japanese_order', '')}")
        doc.add_paragraph(f"パスポート番号：{primary_data.get('passport_number', '')}")
        doc.add_paragraph(f"生年月日（例：1982年12月15日）：{primary_data.get('date_of_birth', '')}")

        # Accompanying Family Members (dynamic, 1-7)
        japanese_numbers = {
            1: '一', 2: '二', 3: '三', 4: '四', 
            5: '五', 6: '六', 7: '七'
        }
        
        for i in range(1, 8):  # Check for accompanying1 through accompanying7
            member_key = f"accompanying{i}"
            member_data = data.get(member_key, {})
            
            if member_data:  # Only add section if member data exists
                japanese_num = japanese_numbers.get(i, str(i))
                doc.add_heading(f"同行家族{japanese_num}", level=1)
                doc.add_paragraph(f"名前：{member_data.get('full_name_japanese_order', '')}")
                doc.add_paragraph(f"パスポート番号：{member_data.get('passport_number', '')}")
                doc.add_paragraph(f"生年月日：{member_data.get('date_of_birth', '')}")

        doc.add_paragraph("※ 全員分のパスポートのコピーを添付しております。")
        doc.add_paragraph("お忙しいところ恐れ入りますが、至急ご対応いただけますと幸いです。")
        doc.add_paragraph("どうぞよろしくお願い申し上げます。")
        
        doc.save(output_path)
        return output_path

    def generate_email_content(self, data: Dict[str, Dict[str, str]]) -> Dict[str, str]:
        """Generate email content for Outlook draft"""
        # Get primary applicant name for subject
        primary_data = data.get("primary", {})
        primary_name = primary_data.get('full_name_japanese_order', 'Applicant')
        
        # Create subject
        subject = f"ビザ申請書類提出 - {primary_name}"
        
        # Build email body
        body_lines = [
            "ご担当者様",
            "",
            "ビザ申請に必要な書類を提出させていただきます。",
            "",
            "重複プロファイル解消のお願い",
            "ビザ面接の予約をしようとしたところ、以下のメッセージが表示されました。",
            "「あなたの個人情報は、当社のデータベースに既に存在するプロファイルと一致します。」",
            "しかし、古いプロファイルのログイン情報がわからず、現在予約手続きを進められない状況です。",
            "お手数ですが、古いプロファイルを削除いただけますようお願い申し上げます。",
            "",
            "【申請者情報】",
            ""
        ]
        
        # Add primary applicant
        if primary_data:
            body_lines.extend([
                "■ 主たる申請者",
                f"名前：{primary_data.get('full_name_japanese_order', '')}",
                f"パスポート番号：{primary_data.get('passport_number', '')}",
                f"生年月日：{primary_data.get('date_of_birth', '')}",
                ""
            ])
        
        # Add family members
        japanese_numbers = {
            1: '一', 2: '二', 3: '三', 4: '四', 
            5: '五', 6: '六', 7: '七'
        }
        
        for i in range(1, 8):
            member_key = f"accompanying{i}"
            member_data = data.get(member_key, {})
            
            if member_data:
                japanese_num = japanese_numbers.get(i, str(i))
                body_lines.extend([
                    f"■ 同行家族{japanese_num}",
                    f"名前：{member_data.get('full_name_japanese_order', '')}",
                    f"パスポート番号：{member_data.get('passport_number', '')}",
                    f"生年月日：{member_data.get('date_of_birth', '')}",
                    ""
                ])
        
        body_lines.extend([
            "※ 全員分のパスポートのコピーを添付してください。",
            "",
            "お忙しいところ恐れ入りますが、至急ご対応いただけますと幸いです。",
            "どうぞよろしくお願い申し上げます。"
        ])
        
        body = "\n".join(body_lines)
        
        return {
            "subject": subject,
            "body": body
        }

    def process_multiple_members(self, members_data: Dict[str, Dict[str, Any]]) -> Dict[str, str]:
        """Process multiple members and generate combined forms"""
        print(f"Processing {len(members_data)} members")
        
        # Map each member's passport data
        processed_data = {}
        
        for member_type, passport_data in members_data.items():
            if passport_data:
                form_data = self.map_passport_to_form_fields(passport_data)
                processed_data[member_type] = form_data
                print(f"Processed {member_type}: {form_data.get('last_name', '')} {form_data.get('first_name', '')}")
        
        # Generate output files using current directory structure
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        os.makedirs('generated_forms', exist_ok=True)
        
        pdf_output = f"generated_forms/japanese_visa_form_{timestamp}.pdf"
        word_output = f"generated_forms/japanese_visa_form_{timestamp}.docx"
        
        pdf_path = self.generate_pdf_form(processed_data, pdf_output)
        word_path = self.generate_word_document(processed_data, word_output)
        
        # Generate email content
        email_content = self.generate_email_content(processed_data)
        
        return {
            "pdf_path": pdf_path,
            "word_path": word_path,
            "email_subject": email_content["subject"],
            "email_body": email_content["body"],
            "success": True,
            "members_processed": len(processed_data)
        }


# Utility functions for integration with main app
japanese_generator = JapaneseFormGenerator()

def generate_japanese_forms_from_passports(members_data: Dict[str, Dict[str, Any]]) -> Dict[str, str]:
    """Main interface function for the web app"""
    return japanese_generator.process_multiple_members(members_data)